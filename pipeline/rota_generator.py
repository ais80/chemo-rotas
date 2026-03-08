"""Rota (H-ROTA style) Word document generator.

Builds a prescription chart from scratch with:
- Header area with regimen name, patient placeholders, clinical notes
- Drug tables: Day, Drug, Calculation, Dose, IV Fluid, Volume, Flow Rate,
  Special Directions, Admin signatures, Time
- Signature rows for Prescriber / Chemo Nurse / Pharmacist
- Additional therapy section (non-sequenced drugs)
- Administration advice from precautions
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from models import (
    ParsedProtocol, ChemoRow, NonSequencedRow, RowType,
)
from utils import is_fluid_row, split_drug_diluent, split_dose_calc_volume


DRUG_TABLE_HEADERS = [
    'DAY NO.\nDATE', 'DRUG', 'CALCULATION', 'DOSE',
    'I.V. FLUID', 'VOLUME\n(ml)', 'FLOW\nRATE',
    'SPECIAL DIRECTIONS',
    'DRUG\nADMINISTRATION\nsig.', 'sig.', 'TIME'
]

DRUG_TABLE_WIDTHS_CM = [1.8, 3.5, 2.5, 2.0, 2.5, 1.5, 2.0, 4.5, 2.0, 2.0, 1.5]


class RotaGenerator:
    def generate(self, protocol: ParsedProtocol, output_path: str):
        doc = Document()

        # Configure page: landscape A4
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

        # 1. Header area
        self._add_header(doc, protocol)

        # 2. Clinical notes / Points to note
        self._add_clinical_notes(doc, protocol)

        # 3. Drug table(s) for IV chemotherapy
        if protocol.treatment.has_chemo_table:
            self._add_drug_table(doc, protocol.treatment.chemo_rows)
        elif protocol.treatment.non_sequenced_rows:
            # Oral-only protocol — add a simplified table
            self._add_oral_drug_table(doc, protocol.treatment.non_sequenced_rows)

        # 4. Signature row
        self._add_signature_table(doc)

        # 5. Additional therapy (non-sequenced / support meds)
        if protocol.treatment.has_chemo_table and protocol.treatment.non_sequenced_rows:
            self._add_additional_therapy(doc, protocol)

        # 6. Administration advice (from precautions)
        if not protocol.precautions.is_empty:
            self._add_admin_advice(doc, protocol)

        doc.save(output_path)

    # ── Header ─────────────────────────────────────────────────────────

    def _add_header(self, doc: Document, protocol: ParsedProtocol):
        """Add the document header with regimen name and patient fields."""
        # Title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(protocol.header.regimen_name)
        title_run.bold = True
        title_run.font.size = Pt(16)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Rota code
        code_para = doc.add_paragraph()
        code_run = code_para.add_run(protocol.header.rota_code)
        code_run.font.size = Pt(10)
        code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Patient details placeholder
        details = doc.add_paragraph()
        details.add_run('\n')
        for label in ['Patient Name:', 'Hospital No:', 'DOB:', 'Consultant:',
                       'Height:', 'Weight:', 'BSA:']:
            details.add_run(f'{label}\t\t\t\t').font.size = Pt(10)
            details.add_run('\n')

        # Cycle and date
        cycle_para = doc.add_paragraph()
        cycle_para.add_run('Cycle No:\t…………………………….\t\t').font.size = Pt(10)
        cycle_para.add_run('Start Date:\t…………………………….').font.size = Pt(10)

    # ── Clinical Notes ─────────────────────────────────────────────────

    def _add_clinical_notes(self, doc: Document, protocol: ParsedProtocol):
        """Add clinical notes from the warnings/tests section."""
        if not protocol.warnings_text:
            return

        doc.add_paragraph()  # Spacer
        notes_para = doc.add_paragraph()
        notes_run = notes_para.add_run('Points to note:')
        notes_run.bold = True
        notes_run.font.size = Pt(10)

        for text in protocol.warnings_text:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(9)

    # ── Drug Table (IV Chemotherapy) ───────────────────────────────────

    def _add_drug_table(self, doc: Document, chemo_rows: list[ChemoRow]):
        """Create the main drug table for IV chemotherapy."""
        # Filter to normal rows only (dose reductions noted separately)
        normal_rows = [r for r in chemo_rows if r.row_type == RowType.NORMAL]
        if not normal_rows:
            return

        doc.add_paragraph()  # Spacer

        table = doc.add_table(rows=1, cols=len(DRUG_TABLE_HEADERS))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_row = table.rows[0]
        for i, header_text in enumerate(DRUG_TABLE_HEADERS):
            cell = header_row.cells[i]
            cell.text = header_text
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        # Set column widths
        for i, width in enumerate(DRUG_TABLE_WIDTHS_CM):
            for row in table.rows:
                row.cells[i].width = Cm(width)

        # Data rows
        for row_data in normal_rows:
            drug_name, diluent = split_drug_diluent(row_data.drug_diluent)
            calc, dose, volume = split_dose_calc_volume(row_data.dose_calc_volume)

            data_row = table.add_row()
            values = [
                row_data.stage_day,         # DAY NO.
                drug_name or diluent,       # DRUG
                calc,                       # CALCULATION
                dose,                       # DOSE
                diluent if drug_name else '',  # I.V. FLUID
                volume,                     # VOLUME
                row_data.rate,              # FLOW RATE
                row_data.special_directions, # SPECIAL DIRECTIONS
                '',                         # Admin sig 1
                '',                         # Admin sig 2
                row_data.time,              # TIME
            ]

            for i, val in enumerate(values):
                if val and i < len(data_row.cells):
                    data_row.cells[i].text = val
                    for para in data_row.cells[i].paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)

    # ── Oral Drug Table ────────────────────────────────────────────────

    def _add_oral_drug_table(self, doc: Document, ns_rows: list[NonSequencedRow]):
        """Create a simplified table for oral-only protocols."""
        normal_rows = [r for r in ns_rows if r.row_type == RowType.NORMAL]
        if not normal_rows:
            return

        doc.add_paragraph()  # Spacer
        headers = ['DRUG', 'DOSE', 'FREQUENCY', 'ROUTE', 'DAYS',
                   'SPECIAL DIRECTIONS', 'DRUG\nADMINISTRATION\nsig.', 'sig.']

        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(8)

        # Data rows
        for row_data in normal_rows:
            days = f"Day {row_data.first_dose_day}"
            if row_data.final_dose_day:
                days += f" - {row_data.final_dose_day}"

            data_row = table.add_row()
            values = [
                row_data.drug,
                row_data.dose_calculation,
                row_data.freq,
                row_data.route,
                days,
                row_data.timing_constraints,
                '', '',  # Signature columns
            ]
            for i, val in enumerate(values):
                if val and i < len(data_row.cells):
                    data_row.cells[i].text = val
                    for para in data_row.cells[i].paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)

    # ── Signature Table ────────────────────────────────────────────────

    def _add_signature_table(self, doc: Document):
        """Add the prescriber/nurse/pharmacist signature row."""
        doc.add_paragraph()  # Spacer

        table = doc.add_table(rows=1, cols=12)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        sig_labels = [
            'Prescriber Sig.', '……………………', 'Date:', '……………',
            'Chemo Nurse Sig.', '…………………….', 'Date:', '………',
            'Pharmacist Sig.', '…………………….', 'Date:', '………',
        ]

        for i, label in enumerate(sig_labels):
            cell = table.rows[0].cells[i]
            cell.text = label
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)

    # ── Additional Therapy ─────────────────────────────────────────────

    def _add_additional_therapy(self, doc: Document, protocol: ParsedProtocol):
        """Add non-sequenced / support medications as an additional therapy section."""
        normal_ns = [r for r in protocol.treatment.non_sequenced_rows
                     if r.row_type == RowType.NORMAL]
        if not normal_ns:
            return

        doc.add_paragraph()  # Spacer
        header_para = doc.add_paragraph()
        header_run = header_para.add_run('Additional Therapy')
        header_run.bold = True
        header_run.font.size = Pt(11)

        for row in normal_ns:
            freq_str = f" {row.freq}" if row.freq else ""
            route_str = f" {row.route}" if row.route else ""
            timing_str = f" ({row.timing_constraints})" if row.timing_constraints else ""
            days_str = ""
            if row.first_dose_day:
                days_str = f" from day {row.first_dose_day}"
                if row.final_dose_day:
                    days_str += f" to day {row.final_dose_day}"

            text = f"{row.drug} {row.dose_calculation}{freq_str}{route_str}{timing_str}{days_str}"
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(9)

    # ── Administration Advice ──────────────────────────────────────────

    def _add_admin_advice(self, doc: Document, protocol: ParsedProtocol):
        """Add administration advice from precautions section."""
        doc.add_paragraph()  # Spacer
        header_para = doc.add_paragraph()
        header_run = header_para.add_run('Administration Advice / Precautions')
        header_run.bold = True
        header_run.font.size = Pt(11)

        for text in protocol.precautions.text_paragraphs:
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.size = Pt(9)
