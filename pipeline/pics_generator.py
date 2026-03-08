"""PICS Word document generator.

Clones the FLAG PICS template and populates it with parsed protocol data.
The template has 4 tables:
  Table 0: CHEMOTHERAPY (sequenced drugs) - 13 cols
  Table 1: FLUIDS - 12 cols
  Table 2: NON-SEQUENCED (oral/support meds) - 13 cols
  Table 3: PROCEED RULES - 5 cols
And paragraph sections for Warnings and Rota Information.
"""

from copy import deepcopy
from lxml import etree
from docx import Document
from docx.shared import Pt
from models import (
    ParsedProtocol, ChemoRow, NonSequencedRow, ProceedRuleDrug, RowType,
)
from utils import is_fluid_row

# XML namespace for Word documents
W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


class PICSGenerator:
    def __init__(self, template_path: str):
        self.template_path = template_path

    def generate(self, protocol: ParsedProtocol, output_path: str):
        doc = Document(self.template_path)

        tables = doc.tables
        if len(tables) < 4:
            raise ValueError(f"Template has {len(tables)} tables, expected at least 4")

        # Separate chemo rows into drugs vs fluids
        drug_rows = [r for r in protocol.treatment.chemo_rows if not is_fluid_row(r.drug_diluent)]
        fluid_rows = [r for r in protocol.treatment.chemo_rows if is_fluid_row(r.drug_diluent)]

        # Table 0: CHEMOTHERAPY
        self._fill_chemo_table(tables[0], drug_rows)

        # Table 1: FLUIDS
        self._fill_fluids_table(tables[1], fluid_rows)

        # Table 2: NON-SEQUENCED
        self._fill_nonsequenced_table(tables[2], protocol.treatment.non_sequenced_rows)

        # Table 3: PROCEED RULES
        self._fill_proceed_rules_table(tables[3], protocol.proceed_rules)

        # Fill paragraph sections
        self._fill_paragraphs(doc, protocol)

        doc.save(output_path)

    # ── Table Population ───────────────────────────────────────────────

    def _fill_chemo_table(self, table, rows: list[ChemoRow]):
        """Populate the CHEMOTHERAPY table (Table 0)."""
        self._clear_data_rows(table, header_rows=2)

        for row_data in rows:
            new_row = self._add_row(table)
            cells = new_row.cells

            cell_values = [
                row_data.stage_day,
                row_data.time,
                row_data.drug_diluent,
                row_data.round_dose_to,
                row_data.dose_calc_volume,
                row_data.rate,
                row_data.route,
                row_data.special_directions,
                row_data.target_interval,
                row_data.margin,
                row_data.follows_seq_label,
                row_data.line,
                row_data.seq_label,
            ]

            for i, val in enumerate(cell_values):
                if i < len(cells):
                    self._set_cell_text(cells[i], val, row_data.row_type)

    def _fill_fluids_table(self, table, rows: list[ChemoRow]):
        """Populate the FLUIDS table (Table 1). Same data as chemo but for fluid rows."""
        self._clear_data_rows(table, header_rows=2)

        for row_data in rows:
            new_row = self._add_row(table)
            cells = new_row.cells

            # Fluids table has 12 cols (no "Round dose to nearest" column)
            cell_values = [
                row_data.stage_day,
                row_data.time,
                row_data.drug_diluent,
                row_data.dose_calc_volume,
                row_data.rate,
                row_data.route,
                row_data.special_directions,
                row_data.target_interval,
                row_data.margin,
                row_data.follows_seq_label,
                row_data.line,
                row_data.seq_label,
            ]

            for i, val in enumerate(cell_values):
                if i < len(cells):
                    self._set_cell_text(cells[i], val, row_data.row_type)

    def _fill_nonsequenced_table(self, table, rows: list[NonSequencedRow]):
        """Populate the NON-SEQUENCED table (Table 2)."""
        self._clear_data_rows(table, header_rows=2)

        for row_data in rows:
            new_row = self._add_row(table)
            cells = new_row.cells

            cell_values = [
                row_data.drug,
                row_data.dose_calculation,
                row_data.mode,
                row_data.freq,
                row_data.timing_constraints,
                row_data.route,
                row_data.form,
                row_data.start_with_oof,
                row_data.first_dose_day,
                row_data.first_dose_time,
                row_data.final_dose_day,
                row_data.final_dose_time,
                row_data.group,
            ]

            for i, val in enumerate(cell_values):
                if i < len(cells):
                    self._set_cell_text(cells[i], val, row_data.row_type)

    def _fill_proceed_rules_table(self, table, rules: list[ProceedRuleDrug]):
        """Populate the PROCEED RULES table (Table 3)."""
        self._clear_data_rows(table, header_rows=1)

        for rule in rules:
            new_row = self._add_row(table)
            cells = new_row.cells

            cell_values = [
                rule.drug_name,
                rule.neutrophils,
                rule.platelets,
                rule.renal,
                rule.hepatic,
            ]

            for i, val in enumerate(cell_values):
                if i < len(cells):
                    self._set_cell_text(cells[i], val)

    # ── Paragraph Sections ─────────────────────────────────────────────

    def _fill_paragraphs(self, doc: Document, protocol: ParsedProtocol):
        """Fill the Warnings and Rota Information paragraph sections."""
        # Find key paragraphs by their text content
        warnings_idx = None
        info_idx = None

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            if text == 'warnings':
                warnings_idx = i
            elif text in ('rota information', 'information'):
                info_idx = i

        # Insert warnings text after the "Warnings" header
        if warnings_idx is not None and protocol.warnings_text:
            self._insert_paragraphs_after(doc, warnings_idx, protocol.warnings_text)

        # Insert info text after the "Rota Information" / "Information" header
        if info_idx is not None and protocol.info_text:
            # Adjust index if we inserted warnings above
            offset = len(protocol.warnings_text) if warnings_idx is not None else 0
            self._insert_paragraphs_after(doc, info_idx + offset, protocol.info_text)

        # If no existing Warnings header, create one
        if warnings_idx is None and protocol.warnings_text:
            doc.add_paragraph('Warnings', style='Normal')
            for text in protocol.warnings_text:
                p = doc.add_paragraph(text, style='Normal')
                for run in p.runs:
                    run.font.size = Pt(9)

        # If no existing Info header, create one
        if info_idx is None and protocol.info_text:
            doc.add_paragraph('Rota Information', style='Normal')
            for text in protocol.info_text:
                p = doc.add_paragraph(text, style='Normal')
                for run in p.runs:
                    run.font.size = Pt(9)

    def _insert_paragraphs_after(self, doc: Document, index: int, texts: list[str]):
        """Insert paragraphs after a given paragraph index."""
        # python-docx doesn't have direct "insert after" support
        # We manipulate the XML directly
        ref_para = doc.paragraphs[index]._element
        parent = ref_para.getparent()

        for text in reversed(texts):
            new_p = deepcopy(ref_para)
            # Clear existing text
            for r in new_p.findall(f'{W_NS}r'):
                new_p.remove(r)
            # Add new run with text
            new_r = etree.SubElement(new_p, f'{W_NS}r')
            new_t = etree.SubElement(new_r, f'{W_NS}t')
            new_t.text = text
            new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            # Set font size to 9pt
            rpr = new_r.find(f'{W_NS}rPr')
            if rpr is None:
                rpr = etree.SubElement(new_r, f'{W_NS}rPr')
            sz = rpr.find(f'{W_NS}sz')
            if sz is None:
                sz = etree.SubElement(rpr, f'{W_NS}sz')
            sz.set(f'{W_NS}val', '18')  # 18 half-points = 9pt
            # Remove bold if present (warnings text is not bold)
            for b in rpr.findall(f'{W_NS}b'):
                rpr.remove(b)
            # Insert after reference
            ref_para.addnext(new_p)

    # ── Row Management ─────────────────────────────────────────────────

    def _clear_data_rows(self, table, header_rows: int = 2):
        """Remove all data rows from a table, keeping header rows."""
        tbl = table._tbl
        rows = tbl.findall(f'{W_NS}tr')
        for row in rows[header_rows:]:
            tbl.remove(row)

    def _add_row(self, table):
        """Add a new row with proper independent cells (no vertical merges)."""
        tbl = table._tbl
        rows = tbl.findall(f'{W_NS}tr')
        template_row = rows[-1]
        new_row = deepcopy(template_row)

        # For every cell in the new row:
        # 1. Remove vMerge elements (which link cells to merged headers)
        # 2. Clear all text content
        for tc in new_row.findall(f'{W_NS}tc'):
            tc_pr = tc.find(f'{W_NS}tcPr')
            if tc_pr is not None:
                for vm in tc_pr.findall(f'{W_NS}vMerge'):
                    tc_pr.remove(vm)
                # Also remove gridSpan > 1 to avoid horizontal merges
                for gs in tc_pr.findall(f'{W_NS}gridSpan'):
                    tc_pr.remove(gs)
            # Clear all runs in all paragraphs
            for p in tc.findall(f'.//{W_NS}p'):
                for r in p.findall(f'{W_NS}r'):
                    p.remove(r)

        tbl.append(new_row)
        return table.rows[-1]

    def _set_cell_text(self, cell, text: str, row_type: RowType = RowType.NORMAL):
        """Set a cell's text content, optionally with strikethrough for dose reductions."""
        if not text:
            return

        # Use the first paragraph in the cell
        if cell.paragraphs:
            para = cell.paragraphs[0]
            # Clear any existing runs first
            p_elem = para._element
            for r in p_elem.findall(f'{W_NS}r'):
                p_elem.remove(r)
            # Add fresh run
            run = para.add_run(text)
            run.font.size = Pt(9)
            run.font.name = 'Arial'

            if row_type == RowType.DOSE_REDUCTION:
                run.font.strike = True
