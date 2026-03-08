"""Reverse parser: Rota .docx -> ParsedProtocol.

Parses H-ROTA style Word documents into the universal intermediate format
for conversion to PICS .docx or HTML .htm files.
"""

import os
import re
from docx import Document
from lxml import etree
from models import (
    ParsedProtocol, ProtocolHeader, ChemoRow, NonSequencedRow,
    ProceedRuleDrug, SectionContent, TreatmentData, RowType,
)
from review_flags import ReviewFlagCollector
from utils import is_fluid_row

W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
VML_NS = '{urn:schemas-microsoft-com:vml}'
WP_NS = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'


class RotaDocxParser:
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        self.doc = Document(docx_path)
        self.flags = ReviewFlagCollector()
        self.filename = os.path.basename(docx_path)

    def parse(self) -> ParsedProtocol:
        """Main entry point. Returns a populated ParsedProtocol."""
        header = self._parse_header()

        # Classify all tables
        classified = self._classify_tables()

        # Parse drug tables -> ChemoRow list
        chemo_rows = self._parse_drug_tables(classified['drug'])

        # Parse supportive meds tables -> NonSequencedRow list
        non_seq_rows = self._parse_supportive_meds(classified['supportive'])

        # Parse dose modification tables -> ProceedRuleDrug list + SectionContent
        proceed_rules, dose_mod_content = self._parse_dose_modifications(
            classified['dose_mod']
        )

        # Parse body paragraphs -> warnings, precautions, info, additional therapy
        warnings, precautions, info_text, additional_ns = self._parse_body_paragraphs()

        # Merge additional therapy from paragraphs into non_seq_rows
        non_seq_rows.extend(additional_ns)

        treatment = TreatmentData(
            chemo_rows=chemo_rows,
            non_sequenced_rows=non_seq_rows,
            has_chemo_table=len(chemo_rows) > 0,
            has_nonsequenced_table=len(non_seq_rows) > 0,
        )

        # Flag sections that Rotas don't contain
        for section_name in ['Eligibility', 'Exclusions', 'Premedications',
                             'Continue treatment if']:
            self.flags.info("missing_section", section_name,
                            f"{section_name} section not available in Rota documents"
                            " - requires manual entry")

        return ParsedProtocol(
            header=header,
            treatment=treatment,
            proceed_rules=proceed_rules,
            dose_modifications=dose_mod_content,
            precautions=SectionContent(
                section_name="Precautions",
                text_paragraphs=precautions,
                is_empty=len(precautions) == 0,
            ),
            support_medications=SectionContent(
                section_name="Support Medications",
                text_paragraphs=[
                    f"{r.drug} {r.dose_calculation}".strip()
                    for r in non_seq_rows
                ] if non_seq_rows else [],
                is_empty=len(non_seq_rows) == 0,
            ),
            warnings_text=warnings,
            info_text=info_text,
            review_flags=self.flags.get_all(),
        )

    # ── Header Extraction ─────────────────────────────────────────────

    def _parse_header(self) -> ProtocolHeader:
        """Extract regimen name, rota code, and last updated from the document."""
        rota_code = ""
        regimen_name = ""
        last_updated = ""

        # 1. Search document header tables for rota code, date, and regimen name
        regimen_candidates = []
        for section in self.doc.sections:
            hdr = section.header
            if hdr is None:
                continue
            for table in hdr.tables:
                for row in table.rows:
                    # Use deduped cells to handle merges
                    deduped = self._dedup_cells(row)
                    cells_text = [c.text.strip() for c in deduped]
                    for i, cell_text in enumerate(cells_text):
                        low = cell_text.lower()
                        # Find rota code
                        if re.search(r'doc(ument)?\s*code', low) and not rota_code:
                            code_match = re.search(
                                r'(?:doc(?:ument)?\s*code[:\s]*)(H-?ROTA?\s*\d+\w*)',
                                cell_text, re.IGNORECASE
                            )
                            if code_match:
                                rota_code = code_match.group(1).strip()
                            elif i + 1 < len(cells_text):
                                rota_code = cells_text[i + 1].strip()
                        # Find issue date
                        if re.search(r'issue\s*date|last\s*updated', low) \
                                and not last_updated:
                            if i + 1 < len(cells_text) and cells_text[i + 1]:
                                last_updated = cells_text[i + 1].strip()
                        # Collect regimen name candidates from large cells
                        if len(cell_text) > 15:
                            # Skip metadata cells
                            skip_kw = [
                                'document code', 'issue date', 'page',
                                'written by', 'ward', 'authorised', 'version',
                                'review date', 'valid until', 'clinical nurse',
                                'pharmacist', 'consultant', 'patient label',
                                'hb', 'na+', 'height', 'weight', 'alb',
                                'recorded by', 'mtx', 'plasma',
                                'concentration', 'ref:', 'spc',
                                'start date', 'cdf number', 'nice',
                            ]
                            if not any(kw in low for kw in skip_kw):
                                # Also skip cells in rows with metadata labels
                                row_text = ' '.join(
                                    c.text.strip().lower()
                                    for c in self._dedup_cells(row)
                                )
                                row_skip = [
                                    'consultant', 'written by', 'authorised',
                                    'clinical nurse', 'pharmacist',
                                ]
                                if any(kw in row_text for kw in row_skip):
                                    continue
                                # Clean: take lines that look like a regimen
                                lines = [
                                    l.strip() for l in cell_text.split('\n')
                                    if l.strip()
                                    and len(l.strip()) > 5
                                    and not re.match(
                                        r'^(H-?ROTA|HROTA|Doc|Page|Written'
                                        r'|Ward|Queen|University|NHS|Centre'
                                        r'|CHEMOTHERAPY$|PRESCRIPTION$'
                                        r'|Ref:|SPC )',
                                        l.strip(), re.IGNORECASE
                                    )
                                ]
                                regimen_candidates.extend(lines)

        # Also check header cells that contain drug/indication text
        # (e.g. H-ROTA 233 has "Blinatumomab\nRelapsed/Refractory ALL\nCycle 1 only")
        for section in self.doc.sections:
            hdr = section.header
            if hdr is None:
                continue
            for table in hdr.tables:
                for row in table.rows:
                    deduped = self._dedup_cells(row)
                    for cell in deduped:
                        text = cell.text.strip()
                        if not text or len(text) < 15:
                            continue
                        low = text.lower()
                        # Skip known metadata
                        if any(kw in low for kw in [
                            'document code', 'patient label', 'hb',
                            'na+', 'start date', 'cdf', 'use h-rota',
                            'written by', 'page', 'version', 'issue',
                            'valid until', 'pharmacist', 'consultant',
                            'queen elizabeth', 'university',
                        ]):
                            continue
                        # Look for multi-line indication cells
                        lines = [l.strip() for l in text.split('\n')
                                 if l.strip()]
                        if len(lines) >= 2:
                            # Clean: remove "Cycle X only" line
                            clean_lines = [
                                l for l in lines
                                if not re.match(
                                    r'Cycle\s+\d+', l, re.IGNORECASE
                                )
                            ]
                            if clean_lines:
                                regimen_candidates.append(
                                    ' - '.join(clean_lines)
                                )

        # Pick best regimen name candidate (longest, filtering generic titles)
        generic = {
            'chemotherapy', 'chemotherapy prescription',
            'haematology chemotherapy',
        }
        for cand in sorted(regimen_candidates, key=len, reverse=True):
            if cand.lower() in generic:
                continue
            regimen_name = cand
            break

        # 2. Search VML textboxes in the document body for regimen name
        if not regimen_name:
            regimen_name = self._search_textboxes_for_name()

        # 3. Search body paragraphs for indication/regimen-like text
        if not regimen_name:
            for para in self.doc.paragraphs[:20]:
                text = para.text.strip()
                if not text or len(text) < 10:
                    continue
                low = text.lower()
                # Skip placeholders
                if any(kw in low for kw in [
                    'patient name', 'hospital no', 'dob', 'consultant',
                    'height', 'weight', 'bsa', 'cycle no', 'start date',
                    'administration advice', 'do not flush', 'important',
                    '…………',
                ]):
                    continue
                # Extract regimen name from "Refs: ..." line
                refs_match = re.match(
                    r'Refs?:\s*(.+)', text, re.IGNORECASE
                )
                if refs_match:
                    regimen_name = refs_match.group(1).strip()
                    break
                # Extract from "Reference: ..." line
                ref_match = re.match(
                    r'Reference:\s*(.+)', text, re.IGNORECASE
                )
                if ref_match:
                    regimen_name = ref_match.group(1).strip()
                    break
                # Check if bold and long enough to be a title
                has_bold = any(r.bold for r in para.runs if r.text.strip())
                if has_bold and len(text) > 15:
                    regimen_name = text
                    break

        # 4. Fallback: extract from filename
        if not rota_code:
            rota_code = self._extract_code_from_filename()
        if not regimen_name:
            regimen_name = self._extract_name_from_filename()

        # Normalize rota code
        rota_code = re.sub(r'\s+', '', rota_code)

        if not regimen_name:
            self.flags.warn("header", "regimen_name",
                            "Could not extract regimen name")

        return ProtocolHeader(
            regimen_name=regimen_name,
            rota_code=rota_code,
            last_updated=last_updated,
        )

    def _search_textboxes_for_name(self) -> str:
        """Search VML textboxes in document body for the regimen name."""
        LAB_KEYWORDS = {
            'hb', 'na+', 'na', 'alb', 'height', 'wcc', 'k+', 'k', 'bili',
            'weight', 'nts', 'u', 'alkphos', 's.a.', 'plt', 'cr', 'ast',
            'gfr', 'ca', 'mg', 'neuts', 'platelets', 'bsa', 'alt',
            'recorded by', 'date', 'start date', 'cycle',
        }
        LAB_CONTENT_KEYWORDS = [
            'concentration', 'plasma', 'μmol', 'umol', 'level',
            'clearance', 'creatinine', 'bilirubin',
        ]
        SIG_KEYWORDS = ['prescriber', 'pharmacist', 'nurse', 'sig.', '……']

        candidates = []
        body = self.doc.element.body
        # Find all txbxContent elements (VML text boxes)
        for elem in body.iter():
            if 'txbxContent' not in elem.tag:
                continue
            for p in elem.findall(f'.//{W_NS}p'):
                texts = [t.text for t in p.findall(f'.//{W_NS}t') if t.text]
                full_text = ''.join(texts).strip()
                if not full_text or len(full_text) < 5:
                    continue
                low = full_text.lower()
                if low in LAB_KEYWORDS:
                    continue
                if any(kw in low for kw in SIG_KEYWORDS):
                    continue
                if any(kw in low for kw in LAB_CONTENT_KEYWORDS):
                    continue
                if re.match(r'^[\d./\s]+$', full_text):
                    continue
                candidates.append(full_text)

        # Return longest candidate (most likely the regimen name)
        if candidates:
            return max(candidates, key=len)
        return ""

    def _extract_code_from_filename(self) -> str:
        name = os.path.splitext(self.filename)[0]
        match = re.search(r'H-?ROTA?\s*(\d+\w*)', name, re.IGNORECASE)
        if match:
            return f"HROTA{match.group(1)}"
        return name.replace(' ', '_')

    def _extract_name_from_filename(self) -> str:
        name = os.path.splitext(self.filename)[0]
        cleaned = re.sub(r'H-?ROTA?\s*\d+\w*', '', name, flags=re.IGNORECASE)
        cleaned = re.sub(r'\s*-?\s*(final|v\d+|draft)\s*$', '', cleaned,
                         flags=re.IGNORECASE)
        cleaned = cleaned.strip(' -')
        return cleaned if len(cleaned) > 2 else ""

    # ── Table Classification ──────────────────────────────────────────

    def _classify_tables(self) -> dict:
        """Categorize every table in the document."""
        result = {
            'drug': [], 'signature': [], 'supportive': [],
            'dose_mod': [], 'other': [],
        }

        for ti, table in enumerate(self.doc.tables):
            rows = table.rows
            if not rows:
                result['other'].append((ti, table))
                continue

            # Deduplicate cells (merged cells appear multiple times)
            first_row_cells = self._dedup_cells(rows[0])
            header_text = ' '.join(c.text.strip().upper() for c in first_row_cells)

            # Signature table: 1 row, contains PRESCRIBER or PHARMACIST
            if len(rows) == 1:
                if any(kw in header_text for kw in
                       ['PRESCRIBER', 'PHARMACIST', 'CHECKER']):
                    result['signature'].append((ti, table))
                    continue

            # Drug table: header contains DAY/NO. and (DRUG or ELECTROLYTE) and DOSE
            has_day = bool(re.search(r'\bDAY\b|\bNO\.\b|\bDATE\b', header_text))
            has_drug = bool(re.search(r'\bDRUG\b|\bELECTROLYTE\b', header_text))
            has_dose = bool(re.search(r'\bDOSE\b|\bCALCULATION\b', header_text))
            if has_day and has_drug and has_dose:
                result['drug'].append((ti, table))
                continue

            # Check second row headers too (some tables have split headers)
            if len(rows) >= 2 and has_day:
                second_text = ' '.join(
                    c.text.strip().upper()
                    for c in self._dedup_cells(rows[1])
                )
                combined = header_text + ' ' + second_text
                if (re.search(r'\bDRUG\b|\bELECTROLYTE\b', combined)
                        and re.search(r'\bDOSE\b|\bCALCULATION\b', combined)):
                    result['drug'].append((ti, table))
                    continue

            # Dose modification: Toxicity/Grade/Action or Neutrophils/Platelets
            if any(kw in header_text for kw in ['TOXICITY', 'NEUROPATHY']):
                result['dose_mod'].append((ti, table))
                continue
            # Grade* with asterisk
            if (re.search(r'\bGRADE\b', header_text)
                    and re.search(r'\bACTION\b', header_text)):
                result['dose_mod'].append((ti, table))
                continue
            # Check for proceed-rule style table (Drug, Neutrophils, Platelets...)
            if len(first_row_cells) >= 4:
                if any(kw in header_text for kw in ['NEUTROPHIL', 'PLATELET']):
                    result['dose_mod'].append((ti, table))
                    continue

            # Supportive meds: small tables mentioning PICS or ADDITIONAL
            if len(first_row_cells) <= 3:
                all_text = ' '.join(
                    c.text.upper() for r in rows[:3] for c in r.cells
                )
                if 'PICS' in all_text or 'ADDITIONAL' in all_text:
                    result['supportive'].append((ti, table))
                    continue

            # Unclassified — check if it looks like supportive meds
            if len(first_row_cells) <= 3 and len(rows) >= 2:
                sample = ' '.join(
                    c.text for r in rows[1:3] for c in r.cells
                ).lower()
                if any(kw in sample for kw in [
                    'mg', 'mcg', 'tablet', 'capsule', 'daily', 'bd', 'tds',
                ]):
                    result['supportive'].append((ti, table))
                    continue

            # Check if it's an unusual drug table (no header, but has
            # drug-like content in expected positions — e.g. day number
            # in first cell, drug names, doses)
            if len(rows) >= 2 and len(first_row_cells) >= 5:
                # Check if first cell of first row looks like a day number
                cell0 = first_row_cells[0].text.strip()
                if re.match(r'-?\d+$', cell0):
                    # Looks like day data, not a header — treat as drug table
                    result['drug'].append((ti, table))
                    self.flags.info(
                        "treatment", "headerless_table",
                        f"Table {ti} has no standard header "
                        f"(first cell='{cell0}') — treating as drug table"
                    )
                    continue

            result['other'].append((ti, table))

        return result

    def _dedup_cells(self, row):
        """Return unique cells from a row (python-docx repeats merged cells)."""
        seen = set()
        unique = []
        for cell in row.cells:
            cell_id = id(cell._tc)
            if cell_id not in seen:
                seen.add(cell_id)
                unique.append(cell)
        return unique

    # ── Drug Table Parsing ────────────────────────────────────────────

    def _parse_drug_tables(self, drug_tables: list) -> list[ChemoRow]:
        """Parse all classified drug tables into ChemoRow objects."""
        all_rows = []
        last_drug_name = ""

        for ti, table in drug_tables:
            rows = table.rows
            if len(rows) < 2:
                continue

            # Detect column mapping from header row
            col_map = self._detect_column_mapping(rows[0])

            # Skip header row(s)
            data_start = 1
            if not col_map:
                # No recognizable header — this is a headerless table
                # Use positional defaults (standard 11-col layout)
                col_map = {
                    'day': 0, 'drug': 1, 'calculation': 2, 'dose': 3,
                    'iv_fluid': 4, 'volume': 5, 'rate': 6, 'special': 7,
                }
                data_start = 0  # All rows are data
            elif len(rows) >= 3:
                # Check if row 1 is also a header continuation
                row1_text = ' '.join(
                    c.text.strip().upper()
                    for c in self._dedup_cells(rows[1])
                )
                if re.search(r'\(ml\)|\bDATE\b|\bFLOW\b|\bSIG\b', row1_text):
                    data_start = 2

            for row in rows[data_start:]:
                deduped = self._dedup_cells(row)
                cells = [c.text.strip() for c in deduped]

                # Skip empty rows
                if not any(c for c in cells):
                    continue

                # Skip cross-reference rows (e.g. "prescribe using H-ROTA 114")
                joined = ' '.join(cells).lower()
                if 'h-rota' in joined and 'prescribe' in joined:
                    self.flags.info("treatment", "cross_ref",
                                    f"Skipped cross-reference row: {joined[:80]}")
                    continue

                # Extract fields using column map
                day = self._get_mapped(cells, col_map, 'day')
                drug = self._get_mapped(cells, col_map, 'drug')
                calc = self._get_mapped(cells, col_map, 'calculation')
                dose = self._get_mapped(cells, col_map, 'dose')
                iv_fluid = self._get_mapped(cells, col_map, 'iv_fluid')
                volume = self._get_mapped(cells, col_map, 'volume')
                rate = self._get_mapped(cells, col_map, 'rate')
                special = self._get_mapped(cells, col_map, 'special')
                time_val = self._get_mapped(cells, col_map, 'time')
                route_col = self._get_mapped(cells, col_map, 'route')

                # Carry forward drug name if empty
                if drug:
                    last_drug_name = drug
                elif last_drug_name:
                    drug = last_drug_name

                # Build combined drug_diluent field
                drug_diluent = self._combine_drug_diluent(drug, iv_fluid)

                # Build combined dose_calc_volume field
                dose_calc_volume = self._combine_dose_calc_volume(
                    calc, dose, volume
                )

                # Extract route from rate column if not in dedicated column
                route = route_col
                cleaned_rate = rate
                if not route:
                    route, cleaned_rate = self._extract_route(rate)

                # Detect strikethrough
                row_type = RowType.NORMAL
                for cell_obj in deduped:
                    if self._cell_has_strikethrough(cell_obj):
                        row_type = RowType.DOSE_REDUCTION
                        break

                chemo_row = ChemoRow(
                    stage_day=self._normalize_day(day),
                    time=time_val,
                    drug_diluent=drug_diluent,
                    dose_calc_volume=dose_calc_volume,
                    rate=cleaned_rate,
                    route=route,
                    special_directions=special,
                    row_type=row_type,
                )
                all_rows.append(chemo_row)

        return all_rows

    def _detect_column_mapping(self, header_row) -> dict:
        """Detect which column index maps to which field from header text."""
        col_map = {}
        cells = self._dedup_cells(header_row)

        for i, cell in enumerate(cells):
            text = cell.text.strip().upper()
            words = text.split()

            if any(w in ('DAY', 'NO.', 'DATE') for w in words) \
                    and 'day' not in col_map:
                col_map['day'] = i
            elif any(w in ('DRUG', 'ELECTROLYTE') for w in words) \
                    and 'drug' not in col_map:
                col_map['drug'] = i
            elif 'CALCULATION' in text and 'calculation' not in col_map:
                col_map['calculation'] = i
            elif text == 'DOSE' or (
                'DOSE' in text and 'ADMINISTRATION' not in text
                and 'dose' not in col_map
            ):
                col_map['dose'] = i
            elif any(w in ('I.V.', 'IV') for w in words) \
                    and 'FLUID' in text and 'iv_fluid' not in col_map:
                col_map['iv_fluid'] = i
            elif 'FLUID' in text and 'iv_fluid' not in col_map:
                col_map['iv_fluid'] = i
            elif 'VOLUME' in text and 'volume' not in col_map:
                col_map['volume'] = i
            elif any(w in ('FLOW', 'RATE') for w in words) \
                    and 'rate' not in col_map:
                col_map['rate'] = i
            elif 'ROUTE' in text and 'route' not in col_map:
                col_map['route'] = i
            elif ('SPECIAL' in text or 'DIRECTION' in text) \
                    and 'special' not in col_map:
                col_map['special'] = i
            elif 'TIME' in text and 'ADMINISTRATION' not in text \
                    and 'time' not in col_map:
                col_map['time'] = i
            # Skip ADMINISTRATION/SIG columns (not needed)

        return col_map

    def _get_mapped(self, cells: list[str], col_map: dict,
                    field: str, default: str = '') -> str:
        idx = col_map.get(field)
        if idx is not None and idx < len(cells):
            return cells[idx]
        return default

    def _combine_drug_diluent(self, drug: str, iv_fluid: str) -> str:
        """Combine drug and IV fluid into PICS format: 'DRUG / Diluent'."""
        drug = drug.strip()
        iv_fluid = iv_fluid.strip()
        if drug and iv_fluid:
            return f"{drug} / {iv_fluid}"
        return drug or iv_fluid

    def _combine_dose_calc_volume(self, calc: str, dose: str,
                                  volume: str) -> str:
        """Combine calculation, dose, and volume into PICS format."""
        calc = calc.strip()
        dose = dose.strip()
        volume = volume.strip()

        main = calc or dose
        if not main:
            return volume

        if volume:
            vol = volume if volume.lower().endswith('ml') else f"{volume}ml"
            return f"{main}/{vol}"
        return main

    def _extract_route(self, rate_text: str) -> tuple[str, str]:
        """Extract route from rate field text. Returns (route, cleaned_rate)."""
        if not rate_text:
            return ('', '')

        text = rate_text.strip()
        upper = text.upper()

        # Explicit route keywords
        route_map = [
            (r'\bIV\s+STAT\b', 'IV', ''),
            (r'\bIV\s+BOLUS\b', 'IV', 'bolus'),
            (r'\bS/?C\s+BOLUS\b', 'SC', 'bolus'),
            (r'\bS/?C\b', 'SC', ''),
            (r'\bIM\b', 'IM', ''),
            (r'\bIT\b', 'IT', ''),
            (r'\borally\b', 'PO', ''),
            (r'\b[Pp]\.?[Oo]\.?\b', 'PO', ''),
            (r'\btopical(ly)?\b', 'TOP', ''),
        ]

        for pattern, route, replacement in route_map:
            if re.search(pattern, text, re.IGNORECASE):
                cleaned = re.sub(pattern, replacement, text,
                                 flags=re.IGNORECASE).strip()
                cleaned = re.sub(r'\s{2,}', ' ', cleaned).strip(' ,;')
                return (route, cleaned)

        # If it contains a duration, assume IV
        if re.search(r'\d+\s*(hours?|hrs?|mins?|minutes?)', text,
                      re.IGNORECASE):
            return ('IV', text)

        return ('', text)

    def _normalize_day(self, day_text: str) -> str:
        """Normalize day text to consistent format."""
        if not day_text:
            return ''
        # Remove "Date:" and variants that get merged into day cells
        text = re.sub(r'\s*Date:?\s*', '', day_text).strip()
        # Remove trailing newlines and whitespace
        text = text.strip()
        if not text:
            return ''
        # Already has "Day" prefix
        if re.match(r'day\s', text, re.IGNORECASE):
            return text
        # "DAY1-3" or "DAY 1-3" pattern
        m = re.match(r'DAY\s*(\d+(?:\s*-\s*\d+)?)', text, re.IGNORECASE)
        if m:
            return f"Day {m.group(1).replace(' ', '')}"
        # Plain number or negative number
        if re.match(r'-?\d+$', text):
            return f"Day {text}"
        # Day range like "1-3"
        if re.match(r'\d+\s*-\s*\d+$', text):
            return f"Day {text.replace(' ', '')}"
        return text

    def _cell_has_strikethrough(self, cell) -> bool:
        """Check if any run in the cell has strikethrough formatting."""
        for para in cell.paragraphs:
            for run in para.runs:
                if run.font.strike:
                    return True
                # Also check XML for w:strike
                rpr = run._element.find(f'{W_NS}rPr')
                if rpr is not None:
                    strike = rpr.find(f'{W_NS}strike')
                    if strike is not None:
                        val = strike.get(f'{W_NS}val', 'true')
                        if val != 'false' and val != '0':
                            return True
        return False

    # ── Supportive Meds Parsing ───────────────────────────────────────

    def _parse_supportive_meds(self, supp_tables: list) -> list[NonSequencedRow]:
        """Parse supportive medication tables into NonSequencedRow objects."""
        rows = []
        for ti, table in supp_tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if not text:
                        continue
                    # Skip header-like text
                    low = text.lower()
                    if any(kw in low for kw in [
                        'additional', 'prescribe on pics', 'supportive',
                        'anti-emetic', 'antiemetic', 'medication',
                    ]):
                        continue
                    # Try to parse each line as a medication
                    for line in text.split('\n'):
                        line = line.strip()
                        if line and len(line) > 3:
                            ns_row = self._parse_med_text(line)
                            if ns_row:
                                rows.append(ns_row)
        return rows

    def _parse_med_text(self, text: str) -> NonSequencedRow | None:
        """Parse free-text medication string into NonSequencedRow fields.

        Examples:
            "Metoclopramide 10mg tds prn" -> drug, dose, freq
            "Aciclovir 200mg qds" -> drug, dose, freq
            "Co-trimoxazole 480mg bd Mon-Wed-Fri" -> drug, dose, freq
        """
        if not text or len(text) < 3:
            return None

        # Skip non-medication text
        low = text.lower()
        if any(kw in low for kw in [
            'prescribe', 'see ', 'refer ', 'note:', 'n.b.',
            'please', 'protocol', 'administration',
        ]):
            return None

        # Pattern: Drug_name dose freq [route] [extra]
        match = re.match(
            r'([A-Za-z][\w\s/-]*?)\s+'          # Drug name
            r'(\d[\d.]*\s*'                      # Dose number
            r'(?:mg|mcg|g|units?|mmol|ml|IU'     # Dose unit
            r'|microgram|milligram)?'
            r'(?:/m[2²]|/kg)?)\s*'               # Optional per-BSA/weight
            r'(.*)',                              # Rest (freq, route, etc.)
            text, re.IGNORECASE
        )

        if match:
            drug = match.group(1).strip()
            dose = match.group(2).strip()
            rest = match.group(3).strip().rstrip('.')

            # Extract route from rest
            route = ''
            for rt_pattern, rt_label in [
                (r'\bIV\b', 'IV'), (r'\bPO\b', 'PO'), (r'\borally\b', 'PO'),
                (r'\bSC\b', 'SC'), (r'\bIM\b', 'IM'), (r'\btopical\b', 'TOP'),
            ]:
                if re.search(rt_pattern, rest, re.IGNORECASE):
                    route = rt_label
                    rest = re.sub(rt_pattern, '', rest,
                                  flags=re.IGNORECASE).strip()

            # Extract frequency
            freq = rest.strip(' ,;')

            return NonSequencedRow(
                drug=drug,
                dose_calculation=dose,
                freq=freq,
                route=route,
            )

        # Fallback: treat whole text as drug name
        return NonSequencedRow(drug=text.strip())

    # ── Dose Modification Parsing ─────────────────────────────────────

    def _parse_dose_modifications(self, tables) -> tuple[
        list[ProceedRuleDrug], SectionContent
    ]:
        """Parse dose modification tables into proceed rules and/or text."""
        proceed_rules = []
        dose_mod_paragraphs = []

        for ti, table in tables:
            rows = table.rows
            if len(rows) < 2:
                continue

            first_row_cells = self._dedup_cells(rows[0])
            header_text = ' '.join(c.text.strip().upper() for c in first_row_cells)

            # Format A: Toxicity | Grade | Action (H-ROTA 233 style)
            if 'TOXICITY' in header_text or (
                'GRADE' in header_text and 'ACTION' in header_text
            ):
                for row in rows[1:]:
                    cells = [c.text.strip() for c in self._dedup_cells(row)]
                    if len(cells) >= 3 and any(c for c in cells):
                        line = f"{cells[0]}: Grade {cells[1]} - {cells[2]}"
                        dose_mod_paragraphs.append(line)
                continue

            # Format B: Drug | Neutrophils | Platelets | Renal | Hepatic
            if any(kw in header_text for kw in ['NEUTROPHIL', 'PLATELET']):
                # Find column positions
                col_positions = {}
                for i, cell in enumerate(first_row_cells):
                    t = cell.text.strip().upper()
                    if 'DRUG' in t:
                        col_positions['drug'] = i
                    elif 'NEUTROPHIL' in t:
                        col_positions['neutrophils'] = i
                    elif 'PLATELET' in t:
                        col_positions['platelets'] = i
                    elif 'RENAL' in t or 'GFR' in t:
                        col_positions['renal'] = i
                    elif 'HEPATIC' in t or 'LIVER' in t or 'BILIRUBIN' in t:
                        col_positions['hepatic'] = i

                for row in rows[1:]:
                    cells = [c.text.strip() for c in self._dedup_cells(row)]
                    if not any(c for c in cells):
                        continue
                    drug = cells[col_positions.get('drug', 0)] \
                        if 'drug' in col_positions and \
                        col_positions['drug'] < len(cells) else ''
                    if not drug:
                        continue
                    proceed_rules.append(ProceedRuleDrug(
                        drug_name=drug,
                        neutrophils=cells[col_positions['neutrophils']]
                        if 'neutrophils' in col_positions
                        and col_positions['neutrophils'] < len(cells) else '',
                        platelets=cells[col_positions['platelets']]
                        if 'platelets' in col_positions
                        and col_positions['platelets'] < len(cells) else '',
                        renal=cells[col_positions['renal']]
                        if 'renal' in col_positions
                        and col_positions['renal'] < len(cells) else '',
                        hepatic=cells[col_positions['hepatic']]
                        if 'hepatic' in col_positions
                        and col_positions['hepatic'] < len(cells) else '',
                    ))
                continue

            # Fallback: store as text
            for row in rows[1:]:
                cells = [c.text.strip() for c in self._dedup_cells(row)]
                line = ' | '.join(c for c in cells if c)
                if line:
                    dose_mod_paragraphs.append(line)

        content = SectionContent(
            section_name="Dose Modifications",
            text_paragraphs=dose_mod_paragraphs,
            is_empty=len(dose_mod_paragraphs) == 0,
        )
        return proceed_rules, content

    # ── Body Paragraph Parsing ────────────────────────────────────────

    def _parse_body_paragraphs(self) -> tuple[
        list[str], list[str], list[str], list[NonSequencedRow]
    ]:
        """Parse body paragraphs into warnings, precautions, info, additional.

        Returns: (warnings, precautions, info_text, additional_ns_rows)
        """
        warnings = []
        precautions = []
        info_text = []
        additional_ns = []

        current_section = None
        SKIP_PATTERNS = [
            'patient name', 'hospital no', 'dob:', 'consultant:',
            'height:', 'weight:', 'bsa:', 'cycle no', 'start date',
            '…………', '.....',
        ]

        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            low = text.lower()

            # Skip patient placeholder fields
            if any(kw in low for kw in SKIP_PATTERNS):
                continue

            # Detect section headers
            if re.match(r'points?\s+to\s+note', low):
                current_section = 'notes'
                continue
            if re.match(r'additional\s+therapy|antiemetics?\s+and\s+additional',
                        low):
                current_section = 'additional'
                continue
            if re.match(r'administration\s+advice', low):
                current_section = 'admin_advice'
                continue
            if re.match(r'prescribing\s+advice', low):
                current_section = 'prescribing'
                continue
            if re.match(r'dose\s+adjust', low):
                current_section = 'dose_adj'
                continue
            if re.match(r'other\s+information', low):
                current_section = 'other_info'
                continue

            # Route to section
            if current_section == 'notes':
                warnings.append(text)
            elif current_section == 'additional':
                ns_row = self._parse_med_text(text)
                if ns_row and ns_row.drug:
                    additional_ns.append(ns_row)
                else:
                    info_text.append(text)
            elif current_section in ('admin_advice', 'prescribing', 'dose_adj'):
                precautions.append(text)
            elif current_section == 'other_info':
                info_text.append(text)
            elif current_section is None:
                # Unattached bold paragraphs = potential clinical notes
                has_bold = any(r.bold for r in para.runs if r.text.strip())
                if has_bold and len(text) > 20:
                    warnings.append(text)

        return warnings, precautions, info_text, additional_ns
