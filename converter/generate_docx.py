"""Generate the 4-table DOCX template from config.

Follows the formal specification in CLAUDE.md Section 6.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import RotaConfig, DrugTemplate, BloodTest


def _add_table_row(table, cells_data: list[str]):
    """Add a row to a table with the given cell values."""
    row = table.add_row()
    for i, val in enumerate(cells_data):
        if i < len(row.cells):
            row.cells[i].text = val
    return row



def generate_docx(config: RotaConfig, output_path: str):
    """Generate the 4-table DOCX template."""
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # === Section header: CHEMOTHERAPY ===
    doc.add_paragraph("CHEMOTHERAPY")

    # === Table 0: Chemotherapy Fluids (12 columns) ===
    t0_headers_r0 = [
        "DUE", "DUE", "DRUG/DILUENT", "DOSE CALCULATION/\nVOLUME",
        "RATE", "ROUTE", "Special directions/ PiP(prepared in pharmacy for any MRAPU drugs)",
        "Critical timings", "Critical timings", "Critical timings", "Line", "Seq. label"
    ]
    t0_headers_r1 = [
        "Stage day", "Time", "DRUG/DILUENT", "DOSE CALCULATION/\nVOLUME",
        "RATE", "ROUTE", "Special directions/ PiP(prepared in pharmacy for any MRAPU drugs)",
        "Target interval", "Margin", "Follows  seq. label", "Line", "Seq. label"
    ]
    table0 = doc.add_table(rows=2, cols=12)
    table0.style = 'Table Grid'
    for i, val in enumerate(t0_headers_r0):
        table0.rows[0].cells[i].text = val
    for i, val in enumerate(t0_headers_r1):
        table0.rows[1].cells[i].text = val
    # Populate injectable drug rows from templates where route is IV or SC
    iv_templates = [t for t in config.templates if t.route.upper() in ("IV", "SC")]
    if iv_templates:
        for t in iv_templates:
            rate_str = t.infusion_duration or ""
            _add_table_row(table0, [
                str(t.first_dose_day),   # Stage day
                "",                       # Time
                t.drug_name_upper,        # Drug/Diluent
                f"{t.dose}{t.units}",    # Dose calculation/volume
                rate_str,                 # Rate
                t.route,                  # Route
                t.timing_constraints,     # Special directions
                "",                       # Target interval
                "",                       # Margin
                "",                       # Follows seq. label
                "",                       # Line
                "",                       # Seq. label
            ])
    else:
        # Leave blank rows if no IV drugs
        for _ in range(10):
            _add_table_row(table0, ["1"] + [""] * 11)

    # === Section: FLUIDS ===
    doc.add_paragraph("")
    doc.add_paragraph("FLUIDS")

    # === Table 1: Non-Sequenced Items (12 columns) ===
    t1_headers_r0 = [
        "DUE", "DUE", "DRUG/DILUENT", "UNIT/\nVOLUME",
        "RATE", "ROUTE", "Special directions",
        "Critical timings", "Critical timings", "Critical timings", "Line", "Seq. label"
    ]
    t1_headers_r1 = [
        "Stage day", "Time", "DRUG/DILUENT", "UNIT/\nVOLUME",
        "RATE", "ROUTE", "Special directions",
        "Target interval", "Margin", "Follows  seq. label", "Line", "Seq. label"
    ]
    table1 = doc.add_table(rows=2, cols=12)
    table1.style = 'Table Grid'
    for i, val in enumerate(t1_headers_r0):
        table1.rows[0].cells[i].text = val
    for i, val in enumerate(t1_headers_r1):
        table1.rows[1].cells[i].text = val
    for _ in range(10):
        _add_table_row(table1, [""] * 12)

    # === Section: NON-SEQUENCED ===
    doc.add_paragraph("")
    doc.add_paragraph("NON-SEQUENCED")

    # === Table 2: Drug Templates (13 columns) ===
    t2_headers_r0 = [
        "Drug", "Dose/ Calculation", "Mode", "Freq",
        "Any timing constraints", "Route", "Form", "Start with OOF?",
        "First dose", "First dose", "Final dose", "Final dose", "Group"
    ]
    t2_headers_r1 = [
        "Drug", "Dose/ Calculation", "Mode", "Freq",
        "Any timing constraints", "Route", "Form", "Start with OOF?",
        "Stage day", "Time", "Stage day", "Time", "Group"
    ]
    table2 = doc.add_table(rows=2, cols=13)
    table2.style = 'Table Grid'
    for i, val in enumerate(t2_headers_r0):
        table2.rows[0].cells[i].text = val
    for i, val in enumerate(t2_headers_r1):
        table2.rows[1].cells[i].text = val

    # Add template data rows (oral templates only — IV/SC injectables go in Table 0)
    for t in (t for t in config.templates if t.route.upper() not in ("IV", "SC")):
        row_data = [
            t.drug_name_upper,
            f"{t.dose}{t.units}",
            f"{t.mode} ",     # trailing space matches example
            t.frequency,
            t.timing_constraints,
            f"{t.route} ",    # trailing space matches example
            t.form,
            "-",
            str(t.first_dose_day),
            "",
            f"{t.final_dose_day} ",  # trailing space matches example
            "",
            t.group,
        ]
        _add_table_row(table2, row_data)

    # Add empty padding rows
    for _ in range(2):
        _add_table_row(table2, [""] * 13)

    # === Section: PROCEED RULES ===
    doc.add_paragraph("")
    doc.add_paragraph(" PROCEED RULES")

    # === Table 3: Proceed Rules / Blood Test Warnings (5 columns) ===
    t3_headers = ["Drug", "Neutrophils", "Platelets",
                  "Renal (estimated by Cockroft Gault)", "Hepatic"]
    table3 = doc.add_table(rows=1, cols=5)
    table3.style = 'Table Grid'
    for i, val in enumerate(t3_headers):
        table3.rows[0].cells[i].text = val

    # Build proceed rules row from blood tests
    neuts_text = ""
    plats_text = ""
    renal_text = ""
    hepatic_text = ""

    for bt in config.blood_tests:
        entry = f"{bt.message_text_line1}\n{bt.message_text_line3}"
        if bt.test_code == "NEUTS":
            neuts_text = entry
        elif bt.test_code == "PLATS":
            plats_text = entry
        elif bt.test_code == "GFR":
            renal_text = entry
        elif bt.test_code in ("BILI", "ALT"):
            if hepatic_text:
                hepatic_text += f"\n{entry}"
            else:
                hepatic_text = entry

    row_data = [
        f"{config.drug_full_name}  ",
        neuts_text,
        plats_text,
        renal_text,
        hepatic_text,
    ]
    _add_table_row(table3, row_data)

    # === Free text sections ===
    doc.add_paragraph("")

    # Warnings
    p = doc.add_paragraph("Warnings")
    p.runs[0].bold = True
    for w in config.warnings_paragraphs:
        doc.add_paragraph(w)

    doc.add_paragraph("")

    # Rota Information
    p = doc.add_paragraph("Rota Information")
    p.runs[0].bold = True
    for info in config.rota_info_paragraphs:
        doc.add_paragraph(info)

    doc.save(output_path)
